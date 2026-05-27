#!/usr/bin/env python3
import asyncio
import json
import os
import sys
from datetime import datetime
from pathlib import Path

# Allow imports from mcp_server package when run directly
sys.path.insert(0, str(Path(__file__).parent.parent))

from dotenv import load_dotenv
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import TextContent, Tool

load_dotenv(Path(__file__).parent.parent / ".env")  # dev: project root
load_dotenv()  # production: current working directory

from mcp_server.tools.crawler import check_url_batch, crawl_page
from mcp_server.tools.file_parsers import parse_ahrefs, parse_gsc_data, parse_screaming_frog
from mcp_server.tools.pagespeed import check_pagespeed
from mcp_server.tools.technical_checks import check_robots, check_sitemap
from mcp_server.checklist.technical import get_checklist_as_dict
from mcp_server.checklist.ui import get_ui_checklist_as_dict

app = Server("seo-audit-server")

# In-memory config store for current audit session
_audit_config: dict = {}

OUTPUT_DIR = Path(os.getenv("REPORT_OUTPUT_DIR", "./reports"))


# ── Tool Definitions ──────────────────────────────────────────────────────────

@app.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(
            name="seo_collect_input",
            description=(
                "Lưu cấu hình audit session: domain, brand, language, data sources, "
                "API key, priority groups, include_recommendations."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "domain": {"type": "string", "description": "Domain cần audit (ví dụ: example.com)"},
                    "brand_info": {"type": "string", "description": "Thương hiệu, ngành, sản phẩm/dịch vụ chính"},
                    "audit_purpose": {"type": "string", "description": "Mục đích audit"},
                    "language": {"type": "string", "description": "Ngôn ngữ báo cáo: vi | en"},
                    "pagespeed_api_key": {"type": "string", "description": "Google PageSpeed Insights API key"},
                    "priority_groups": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Các nhóm tiêu chí cần kiểm tra (để trống = tất cả)",
                    },
                    "include_recommendations": {
                        "type": "boolean",
                        "description": "Có sinh đề xuất xử lý cho từng lỗi không",
                    },
                    "data_sources": {
                        "type": "object",
                        "description": "Đường dẫn file dữ liệu bổ sung",
                        "properties": {
                            "screaming_frog": {"type": "string"},
                            "gsc_coverage": {"type": "string"},
                            "gsc_performance": {"type": "string"},
                            "ahrefs": {"type": "string"},
                        },
                    },
                    "output_dir": {
                        "type": "string",
                        "description": "Thư mục lưu báo cáo (để trống = ~/Documents/SEO Audit Reports)",
                        "default": "",
                    },
                    "output_format": {
                        "type": "string",
                        "description": "Định dạng file xuất: md | xlsx | docx",
                        "default": "md",
                    },
                },
                "required": ["domain"],
            },
        ),
        Tool(
            name="seo_crawl_page",
            description=(
                "Crawl một URL và trả về đầy đủ dữ liệu SEO: title, meta, headings, "
                "structured data, links, images, GTM/GA4, header/footer elements, content signals."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL cần crawl (phải có http/https)"},
                    "page_type": {
                        "type": "string",
                        "description": "Loại trang: auto | homepage | product_category | product_detail | blog_category | article | about | contact | 404",
                        "default": "auto",
                    },
                    "follow_redirects": {"type": "boolean", "default": True},
                    "timeout": {"type": "integer", "default": 15, "description": "Timeout tính bằng giây"},
                },
                "required": ["url"],
            },
        ),
        Tool(
            name="seo_check_robots",
            description="Fetch và phân tích robots.txt. Kiểm tra Googlebot allow/disallow, khai báo sitemap, llms.txt.",
            inputSchema={
                "type": "object",
                "properties": {
                    "domain": {"type": "string", "description": "Domain (ví dụ: example.com hoặc https://example.com)"},
                },
                "required": ["domain"],
            },
        ),
        Tool(
            name="seo_check_sitemap",
            description="Fetch và validate sitemap.xml. Đếm URLs, phát hiện sitemap index, kiểm tra lastmod.",
            inputSchema={
                "type": "object",
                "properties": {
                    "domain": {"type": "string", "description": "Domain cần kiểm tra"},
                    "validate_urls": {
                        "type": "boolean",
                        "default": False,
                        "description": "Có check status code cho các URL trong sitemap không (chậm hơn)",
                    },
                    "max_validate": {
                        "type": "integer",
                        "default": 20,
                        "description": "Số URL tối đa để validate",
                    },
                },
                "required": ["domain"],
            },
        ),
        Tool(
            name="seo_check_pagespeed",
            description="Gọi Google PageSpeed Insights API v5. Trả về Core Web Vitals, performance score, top opportunities.",
            inputSchema={
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL cần đo"},
                    "strategy": {
                        "type": "string",
                        "description": "mobile | desktop",
                        "default": "mobile",
                    },
                    "api_key": {
                        "type": "string",
                        "description": "PageSpeed API key (để trống nếu không có, sẽ bị rate limit)",
                        "default": "",
                    },
                },
                "required": ["url"],
            },
        ),
        Tool(
            name="seo_parse_screaming_frog",
            description="Parse file CSV/XLSX export từ Screaming Frog. Phân tích missing titles, duplicate, 404, missing alt text.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Đường dẫn tuyệt đối đến file CSV/XLSX"},
                },
                "required": ["file_path"],
            },
        ),
        Tool(
            name="seo_parse_gsc_data",
            description="Parse file CSV export từ Google Search Console. Hỗ trợ performance và coverage reports.",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Đường dẫn tuyệt đối đến file CSV"},
                    "report_type": {
                        "type": "string",
                        "description": "performance | coverage",
                        "default": "performance",
                    },
                },
                "required": ["file_path"],
            },
        ),
        Tool(
            name="seo_get_checklist",
            description="Trả về toàn bộ checklist: Technical (105+ items, 14 nhóm) + UI (10 page types). Dùng cho Agent 3.",
            inputSchema={
                "type": "object",
                "properties": {
                    "type": {
                        "type": "string",
                        "description": "technical | ui | all",
                        "default": "all",
                    },
                },
            },
        ),
        Tool(
            name="seo_check_url_batch",
            description="Kiểm tra status code của nhiều URL cùng lúc (max 50). Dùng để tìm 404, check www/non-www redirect.",
            inputSchema={
                "type": "object",
                "properties": {
                    "urls": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Danh sách URL cần kiểm tra (tối đa 50)",
                    },
                    "timeout": {
                        "type": "integer",
                        "default": 10,
                        "description": "Timeout mỗi request tính bằng giây",
                    },
                },
                "required": ["urls"],
            },
        ),
        Tool(
            name="seo_save_report",
            description="Render và lưu báo cáo audit. Hỗ trợ định dạng: md (Markdown), xlsx (Excel), docx (Word). Trả về đường dẫn file đã tạo.",
            inputSchema={
                "type": "object",
                "properties": {
                    "audit_results": {
                        "type": "object",
                        "description": "Kết quả phân tích đầy đủ từ Agent 3 (dict với categories, score, top_issues, recommendations)",
                    },
                    "output_dir": {
                        "type": "string",
                        "description": "Thư mục lưu báo cáo (để trống = lấy từ config đã save hoặc ~/Documents/SEO Audit Reports)",
                        "default": "",
                    },
                    "output_format": {
                        "type": "string",
                        "description": "Định dạng file: md | xlsx | docx (để trống = lấy từ config đã save hoặc md)",
                        "default": "",
                    },
                },
                "required": ["audit_results"],
            },
        ),
    ]


# ── Tool Handlers ─────────────────────────────────────────────────────────────

@app.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    try:
        result = await _dispatch(name, arguments)
    except Exception as e:
        result = {"error": str(e), "tool": name}

    return [TextContent(type="text", text=json.dumps(result, ensure_ascii=False, indent=2))]


async def _dispatch(name: str, args: dict) -> dict:
    if name == "seo_collect_input":
        return _handle_collect_input(args)

    elif name == "seo_crawl_page":
        return await crawl_page(
            url=args["url"],
            page_type=args.get("page_type", "auto"),
            follow_redirects=args.get("follow_redirects", True),
            timeout=args.get("timeout", 15),
        )

    elif name == "seo_check_robots":
        return await check_robots(args["domain"])

    elif name == "seo_check_sitemap":
        return await check_sitemap(
            domain=args["domain"],
            validate_urls=args.get("validate_urls", False),
            max_validate=args.get("max_validate", 20),
        )

    elif name == "seo_check_pagespeed":
        api_key = (
            args.get("api_key")
            or _audit_config.get("pagespeed_api_key")
            or os.getenv("PAGESPEED_API_KEY", "")
        )
        return await check_pagespeed(
            url=args["url"],
            strategy=args.get("strategy", "mobile"),
            api_key=api_key,
        )

    elif name == "seo_parse_screaming_frog":
        return parse_screaming_frog(args["file_path"])

    elif name == "seo_parse_gsc_data":
        return parse_gsc_data(
            file_path=args["file_path"],
            report_type=args.get("report_type", "performance"),
        )

    elif name == "seo_get_checklist":
        checklist_type = args.get("type", "all")
        result: dict = {}
        if checklist_type in ("technical", "all"):
            result["technical"] = get_checklist_as_dict()
        if checklist_type in ("ui", "all"):
            result["ui"] = get_ui_checklist_as_dict()
        result["technical_count"] = len(result.get("technical", []))
        result["ui_count"] = len(result.get("ui", []))
        return result

    elif name == "seo_check_url_batch":
        return {
            "results": await check_url_batch(
                urls=args["urls"],
                timeout=args.get("timeout", 10),
            )
        }

    elif name == "seo_save_report":
        return _handle_save_report(args)

    else:
        return {"error": f"Unknown tool: {name}"}


def _handle_collect_input(args: dict) -> dict:
    global _audit_config
    default_output_dir = str(Path.home() / "Documents" / "SEO Audit Reports")
    _audit_config = {
        "domain": args.get("domain", ""),
        "brand_info": args.get("brand_info", ""),
        "audit_purpose": args.get("audit_purpose", "audit toàn diện"),
        "language": args.get("language", "vi"),
        "pagespeed_api_key": args.get("pagespeed_api_key", "") or os.getenv("PAGESPEED_API_KEY", ""),
        "priority_groups": args.get("priority_groups", []),
        "include_recommendations": args.get("include_recommendations", True),
        "data_sources": args.get("data_sources", {}),
        "output_dir": args.get("output_dir", "").strip() or default_output_dir,
        "output_format": args.get("output_format", "md").strip().lower() or "md",
        "created_at": datetime.now().isoformat(),
    }
    return {
        "status": "ok",
        "message": "Cấu hình audit đã được lưu.",
        "config": _audit_config,
    }


def _handle_save_report(args: dict) -> dict:
    audit_results = args.get("audit_results", {})

    fmt_raw = (
        args.get("output_format", "").strip().lower()
        or _audit_config.get("output_format", "md")
    )
    fmt_map = {"excel": "xlsx", "word": "docx", "docs": "docx", "md": "md", "xlsx": "xlsx", "docx": "docx"}
    fmt = fmt_map.get(fmt_raw, "md")

    default_dir = str(Path.home() / "Documents" / "SEO Audit Reports")
    output_dir = Path(
        args.get("output_dir", "").strip()
        or _audit_config.get("output_dir", "")
        or os.getenv("REPORT_OUTPUT_DIR", default_dir)
    )
    output_dir.mkdir(parents=True, exist_ok=True)

    domain = _audit_config.get("domain", "unknown").replace("https://", "").replace("http://", "").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"seo_report_{domain}_{timestamp}.{fmt}"
    output_path = output_dir / filename

    if fmt == "xlsx":
        _save_as_excel(audit_results, output_path)
    elif fmt == "docx":
        _save_as_docx(audit_results, output_path)
    else:
        _save_as_markdown(audit_results, output_path)

    return {
        "status": "ok",
        "file_path": str(output_path.resolve()),
        "filename": filename,
        "format": fmt,
    }


def _save_as_markdown(audit_results: dict, output_path: Path) -> None:
    template_path = Path(__file__).parent / "templates" / "report_template.md"
    if template_path.exists():
        try:
            from jinja2 import Environment, FileSystemLoader
            env = Environment(loader=FileSystemLoader(str(template_path.parent)), autoescape=False)
            content = env.get_template(template_path.name).render(
                config=_audit_config,
                results=audit_results,
                generated_at=datetime.now().strftime("%d/%m/%Y %H:%M"),
            )
        except Exception as e:
            content = _build_fallback_report(audit_results, str(e))
    else:
        content = _build_fallback_report(audit_results)
    output_path.write_text(content, encoding="utf-8")


def _save_as_excel(audit_results: dict, output_path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    STATUS_FILL = {
        "passed":  "C6EFCE",
        "failed":  "FFC7CE",
        "warning": "FFEB9C",
        "manual":  "BDD7EE",
    }
    HEADER_FILL = PatternFill("solid", fgColor="4472C4")
    HEADER_FONT = Font(bold=True, color="FFFFFF")

    def _hdr(ws, headers: list[str]) -> None:
        ws.append(headers)
        for cell in ws[1]:
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL

    def _col_widths(ws, widths: list[int]) -> None:
        for i, w in enumerate(widths, 1):
            from openpyxl.utils import get_column_letter
            ws.column_dimensions[get_column_letter(i)].width = w

    score = audit_results.get("score", {})
    top_issues = audit_results.get("top_issues", [])
    categories = audit_results.get("categories", {})
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    domain = _audit_config.get("domain", "N/A")

    wb = Workbook()

    # ── Sheet 1: Tổng Quan ────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Tổng Quan"
    ws1.append(["SEO Audit Report"])
    ws1["A1"].font = Font(bold=True, size=16)
    ws1.append(["Domain:", domain])
    ws1.append(["Thương hiệu:", _audit_config.get("brand_info", "")])
    ws1.append(["Ngày tạo:", generated_at])
    ws1.append([])
    ws1.append(["Điểm tổng thể:", f"{score.get('percentage', 'N/A')}%  (Grade: {score.get('grade', 'N/A')})"])
    ws1["A6"].font = Font(bold=True)
    ws1.append(["Đạt:", score.get("passed", 0)])
    ws1.append(["Lỗi:", score.get("failed", 0)])
    ws1.append(["Cần cải thiện:", score.get("warning", 0)])
    ws1.append(["Cần kiểm tra thủ công:", score.get("manual", 0)])
    _col_widths(ws1, [30, 45])

    # ── Sheet 2: Top Issues ───────────────────────────────────────────────
    ws2 = wb.create_sheet("Top Issues")
    _hdr(ws2, ["#", "Nhóm", "Tiêu chí", "Mức độ", "Trạng thái", "Ghi nhận", "Đề xuất"])
    for i, issue in enumerate(top_issues[:10], 1):
        status = issue.get("status", "")
        ws2.append([
            i,
            issue.get("category", ""),
            issue.get("name", ""),
            issue.get("priority", ""),
            status,
            issue.get("evidence", ""),
            issue.get("recommendation", ""),
        ])
        color = STATUS_FILL.get(status)
        if color:
            ws2.cell(row=i + 1, column=5).fill = PatternFill("solid", fgColor=color)
    _col_widths(ws2, [5, 25, 45, 15, 15, 50, 50])

    # ── Sheet 3: Chi Tiết ─────────────────────────────────────────────────
    ws3 = wb.create_sheet("Chi Tiết")
    _hdr(ws3, ["Nhóm", "ID", "Tiêu chí", "Mức độ", "Phương pháp", "Trạng thái", "Ghi nhận", "Đề xuất"])
    row_num = 2
    for cat_name, items in categories.items():
        if not isinstance(items, list):
            continue
        for item in items:
            status = item.get("status", "")
            ws3.append([
                cat_name,
                item.get("id", ""),
                item.get("name", ""),
                item.get("priority", ""),
                item.get("check_method", ""),
                status,
                item.get("evidence", ""),
                item.get("recommendation", ""),
            ])
            color = STATUS_FILL.get(status)
            if color:
                ws3.cell(row=row_num, column=6).fill = PatternFill("solid", fgColor=color)
            for col in (7, 8):
                ws3.cell(row=row_num, column=col).alignment = Alignment(wrap_text=True, vertical="top")
            row_num += 1
    _col_widths(ws3, [30, 15, 45, 15, 12, 15, 50, 50])

    wb.save(output_path)


def _save_as_docx(audit_results: dict, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    STATUS_ICONS = {"passed": "✅", "failed": "❌", "warning": "⚠️", "manual": "🔍"}

    score = audit_results.get("score", {})
    top_issues = audit_results.get("top_issues", [])
    categories = audit_results.get("categories", {})
    domain = _audit_config.get("domain", "N/A")
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")

    doc = Document()

    doc.add_heading(f"SEO Audit Report — {domain}", 0)

    p = doc.add_paragraph()
    p.add_run("Ngày tạo: ").bold = True
    p.add_run(generated_at)
    p = doc.add_paragraph()
    p.add_run("Thương hiệu: ").bold = True
    p.add_run(_audit_config.get("brand_info", ""))

    doc.add_heading("Tổng Điểm", 1)
    p = doc.add_paragraph()
    p.add_run(f"Điểm tổng thể: {score.get('percentage', 'N/A')}% — Grade {score.get('grade', 'N/A')}").bold = True
    doc.add_paragraph(
        f"✅ Đạt: {score.get('passed', 0)}  |  "
        f"❌ Lỗi: {score.get('failed', 0)}  |  "
        f"⚠️ Cần cải thiện: {score.get('warning', 0)}  |  "
        f"🔍 Thủ công: {score.get('manual', 0)}"
    )

    doc.add_heading("Top 10 Vấn Đề Ưu Tiên", 1)
    if top_issues:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(["#", "Tiêu chí", "Mức độ", "Trạng thái", "Ghi nhận"]):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for i, issue in enumerate(top_issues[:10], 1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = issue.get("name", "")
            cells[2].text = issue.get("priority", "")
            icon = STATUS_ICONS.get(issue.get("status", ""), "")
            cells[3].text = f"{icon} {issue.get('status', '')}"
            cells[4].text = issue.get("evidence", "")

    doc.add_heading("Chi Tiết Theo Nhóm", 1)
    for cat_name, items in categories.items():
        doc.add_heading(cat_name, 2)
        if isinstance(items, list):
            for item in items:
                icon = STATUS_ICONS.get(item.get("status", ""), "❓")
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{icon} {item.get('name', '')} [{item.get('priority', '')}]").bold = True
                if item.get("evidence"):
                    doc.add_paragraph(f"Ghi nhận: {item['evidence']}", style="List Bullet 2")
                if item.get("recommendation"):
                    doc.add_paragraph(f"Đề xuất: {item['recommendation']}", style="List Bullet 2")

    doc.save(output_path)


def _build_fallback_report(results: dict, template_error: str = "") -> str:
    domain = _audit_config.get("domain", "N/A")
    brand = _audit_config.get("brand_info", "")
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    score = results.get("score", {})
    top_issues = results.get("top_issues", [])
    categories = results.get("categories", {})

    lines = [
        f"# Báo Cáo SEO Audit — {domain}",
        f"\n**Ngày tạo:** {generated_at}",
        f"**Domain:** {domain}",
        f"**Thương hiệu:** {brand}",
        "",
        "---",
        "",
        "## Tổng Điểm",
        f"- **Điểm:** {score.get('percentage', 'N/A')}% (Grade: {score.get('grade', 'N/A')})",
        f"- Đạt: {score.get('passed', 0)} | Lỗi: {score.get('failed', 0)} | Cần cải thiện: {score.get('warning', 0)} | Cần kiểm tra thủ công: {score.get('manual', 0)}",
        "",
        "---",
        "",
        "## Top 10 Vấn Đề Ưu Tiên",
        "",
    ]

    if top_issues:
        for i, issue in enumerate(top_issues[:10], 1):
            lines.append(f"{i}. **[{issue.get('priority', '').upper()}]** {issue.get('name', '')} — {issue.get('status', '')}")
            if issue.get("evidence"):
                lines.append(f"   - Ghi nhận: {issue['evidence']}")
            if issue.get("recommendation"):
                lines.append(f"   - Đề xuất: {issue['recommendation']}")
    else:
        lines.append("_(Không có dữ liệu)_")

    lines += ["", "---", "", "## Chi Tiết Theo Nhóm", ""]

    for cat_name, items in categories.items():
        lines.append(f"### {cat_name}")
        lines.append("")
        if isinstance(items, list):
            for item in items:
                status_icon = {"passed": "✅", "failed": "❌", "warning": "⚠️", "manual": "🔍"}.get(item.get("status", ""), "❓")
                lines.append(f"- {status_icon} **{item.get('name', '')}** ({item.get('priority', '')})")
                if item.get("evidence"):
                    lines.append(f"  > {item['evidence']}")
                if item.get("recommendation"):
                    lines.append(f"  > **Đề xuất:** {item['recommendation']}")
        lines.append("")

    if template_error:
        lines += ["", f"_Ghi chú: Không thể dùng Jinja2 template ({template_error}), dùng fallback renderer._"]

    return "\n".join(lines)


async def main() -> None:
    async with stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            app.create_initialization_options(),
        )


def cli() -> None:
    """Entry point for the `seo-audit-mcp` console script."""
    asyncio.run(main())


if __name__ == "__main__":
    cli()
